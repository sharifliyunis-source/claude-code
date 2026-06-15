from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.2)
    section.right_margin  = Cm(2.2)

# ── Colour palette ────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x0A, 0x1F, 0x44)   # header / table header
GOLD   = RGBColor(0xC8, 0x96, 0x20)   # accent / rule
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
LGRAY  = RGBColor(0xF2, 0xF4, 0xF7)   # alternate row
DGRAY  = RGBColor(0x33, 0x33, 0x33)   # body text
RED_H  = RGBColor(0xC0, 0x39, 0x2B)   # risk red
GREEN_H= RGBColor(0x1A, 0x73, 0x48)   # opportunity green

# ── Helper: set paragraph / run colour ───────────────────────────────────────
def set_cell_bg(cell, rgb):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement('w:shd')
    # RGBColor stores components as tuple indices
    r, g, b = int(rgb[0]), int(rgb[1]), int(rgb[2])
    hex_color = '{:02X}{:02X}{:02X}'.format(r, g, b)
    shd.set(qn('w:val'),   'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'),  hex_color)
    tcPr.append(shd)

def add_bottom_border(paragraph, colour='C89620', size=12):
    pPr  = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot  = OxmlElement('w:bottom')
    bot.set(qn('w:val'),   'single')
    bot.set(qn('w:sz'),    str(size))
    bot.set(qn('w:space'), '1')
    bot.set(qn('w:color'), colour)
    pBdr.append(bot)
    pPr.append(pBdr)

def heading(text, level=1, colour=NAVY, size=None, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT):
    p   = doc.add_paragraph()
    p.alignment = align
    run = p.add_run(text)
    run.bold       = bold
    run.font.color.rgb = colour
    run.font.size  = Pt(size or (20 if level == 1 else 13 if level == 2 else 11))
    run.font.name  = 'Calibri'
    p.paragraph_format.space_before = Pt(10 if level > 1 else 4)
    p.paragraph_format.space_after  = Pt(4)
    return p

def body(text, bold_phrases=None, size=9.5, colour=DGRAY, space_after=4):
    p   = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    # simple bold-marker: wrap **text** segments
    import re
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2])
            r.bold = True
        else:
            r = p.add_run(part)
        r.font.size  = Pt(size)
        r.font.color.rgb = colour
        r.font.name  = 'Calibri'
    return p

def bullet(text, level=0, marker='•'):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent   = Inches(0.25 * (level + 1))
    p.paragraph_format.space_after   = Pt(3)
    p.paragraph_format.space_before  = Pt(1)
    import re
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2])
            r.bold = True
        else:
            r = p.add_run(part)
        r.font.size  = Pt(9.5)
        r.font.color.rgb = DGRAY
        r.font.name  = 'Calibri'
    return p

def section_rule():
    p = doc.add_paragraph()
    add_bottom_border(p)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(6)

def make_table(headers, rows, col_widths=None, alt_rows=True, hdr_bg=NAVY, hdr_fg=WHITE):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # header row
    hdr = t.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, hdr_bg)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p   = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold           = True
        run.font.color.rgb = hdr_fg
        run.font.size      = Pt(8)
        run.font.name      = 'Calibri'

    # data rows
    for ri, row_data in enumerate(rows):
        row = t.rows[ri + 1]
        bg  = LGRAY if (alt_rows and ri % 2 == 0) else WHITE
        for ci, cell_text in enumerate(row_data):
            cell = row.cells[ci]
            set_cell_bg(cell, bg)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.TOP
            p = cell.paragraphs[0]
            import re
            parts = re.split(r'(\*\*[^*]+\*\*)', cell_text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    r = p.add_run(part[2:-2])
                    r.bold = True
                else:
                    r = p.add_run(part)
                r.font.size  = Pt(8)
                r.font.color.rgb = DGRAY
                r.font.name  = 'Calibri'

    # column widths
    if col_widths:
        for ri2, row in enumerate(t.rows):
            for ci2, cell in enumerate(row.cells):
                cell.width = Cm(col_widths[ci2])

    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return t

# ══════════════════════════════════════════════════════════════════════════════
# DOCUMENT CONTENT
# ══════════════════════════════════════════════════════════════════════════════

# ── Cover / Title block ───────────────────────────────────────────────────────
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('EMERGING MARKETS INFRASTRUCTURE INTELLIGENCE BRIEF')
r.bold = True; r.font.size = Pt(18); r.font.color.rgb = NAVY; r.font.name = 'Calibri'

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run('WEEK OF JUNE 15, 2026')
r2.bold = True; r2.font.size = Pt(13); r2.font.color.rgb = GOLD; r2.font.name = 'Calibri'

p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run('Central Asia  ·  South Caucasus  ·  Türkiye  |  Middle Corridor Capital Deployment Monitor')
r3.font.size = Pt(10); r3.font.color.rgb = DGRAY; r3.font.name = 'Calibri'; r3.italic = True

section_rule()

# ── 1. Executive Summary ──────────────────────────────────────────────────────
heading('EXECUTIVE SUMMARY', level=2, colour=NAVY, size=12)
body(
    'Over **USD 12B+** in multilateral development bank capital has been committed across the '
    '**Central Asia–South Caucasus–Türkiye corridor** in the current cycle, anchored by coordinated '
    'sovereign and asset-level investments from **EBRD, World Bank, ADB, AIIB, IsDB, and EDB**. '
    'The dominant theme is **Middle Corridor throughput maximization**: a structurally permanent freight '
    'rerouting away from Russia, now backed by an end-to-end capital stack spanning **Eastern Kazakhstan '
    'to the Bosphorus**. The single most important forward signal is the **June 16–19 convergence of '
    'IsDB Baku Annual Meetings and TIIF Tashkent**, where an estimated **USD 436.67M Karabakh Irrigation '
    'signing** and potential **NDB Uzbekistan debut** will confirm whether diplomatic momentum is converting '
    'to disbursement-grade commitments.'
)
section_rule()

# ── 2. Strategic Developments Portfolio ──────────────────────────────────────
heading('STRATEGIC DEVELOPMENTS PORTFOLIO', level=2, colour=NAVY, size=12)

port_headers = ['Project', 'Institution(s)', 'Value', 'Country', 'Sector', 'Signal']
port_rows = [
    ['INRAIL Istanbul North Rail Crossing',
     '**World Bank** + 5 MDBs', '**USD 6.75B**', '**Türkiye**', 'Rail / Bosphorus freight',
     'Transforms cross-Bosphorus freight capacity from 3M to 50M tonnes/yr — western terminus anchor of the entire Middle Corridor.'],
    ['TC-GATE Georgia Rail Electrification',
     '**World Bank** + **ADB** + **AIIB**', '**USD 729M**', '**Georgia**', 'Rail / corridor',
     'Three-MDB coalition resolves Georgia\'s binding throughput constraint (3–6M t/yr); irreplaceable land bridge node unlocked.'],
    ['KTZ Mointy–Kyzylzhar Greenfield Rail (322 km)',
     '**World Bank** + **AIIB**', '**USD 1.41B**', '**Kazakhstan**', 'Rail',
     'Eastern Kazakhstan corridor spine; value stranded without TC-GATE resolution — now connected.'],
    ['EBRD Regional Support Programme',
     '**EBRD** + **EU**', '**EUR 5B**', '**Azerbaijan**', 'Multi-sector',
     'Largest single EBRD envelope in-window; Karabakh reconstruction + SME + green urban transition simultaneously.'],
    ['EBRD–Kazakhstan EPFA + KTZ Eurobond',
     '**EBRD**', '**EUR 12M grants + USD 125M bond**', '**Kazakhstan**', 'Rail / sovereign policy',
     'Dual-track: sovereign policy alignment (EPFA) + asset-level capital (Eurobond); Middle Corridor institutional anchor.'],
    ['IsDB Kazakhstan Special Economic & Industrial Zones',
     '**IsDB**', '**USD 1.307B**', '**Kazakhstan**', 'Industrial / SEZ',
     'Largest IsDB commitment in Kazakhstan on record; GCC capital as structural co-financier of corridor-adjacent industry.'],
    ['EBRD + IFC Türkiye Industrial Decarbonization (TIDIP)',
     '**EBRD** + **IFC**', '**EUR 5B (2030 target)**', '**Türkiye**', 'Industry / CCUS',
     'First CCUS-linked MDB platform in the region; cement-sector decarbonization establishes bankability template.'],
    ['AIIB Tajikistan Infrastructure Pipeline MOU',
     '**AIIB**', '**~USD 800M**', '**Tajikistan**', 'Multi-sector',
     'Largest single institutional commitment to Tajikistan; pipeline underpopulated — first-mover private co-investment window.'],
    ['World Bank Green & Future Cities',
     '**World Bank (IBRD)**', '**EUR 191.5M**', '**Türkiye**', 'Urban transit / water',
     'Antalya + Konya tramlines + low-emission buses; NDC alignment archetype replicable across 12+ Turkish municipalities.'],
    ['IsDB Azerbaijan Karabakh Irrigation',
     '**IsDB**', '**USD 436.67M**', '**Azerbaijan**', 'Agriculture / reconstruction',
     'Signing expected June 16–19 Baku; post-conflict reconstruction signal to GCC co-investors.'],
    ['EBRD Mirny Wind Farm + BESS (1 GW + 300 MW/600 MWh)',
     '**EBRD**', '**USD 548M**', '**Kazakhstan**', 'Renewable energy',
     'Largest wind-plus-storage transaction in CIS; establishes bankable BESS project template for the region.'],
    ['AIIB Baku Metro Expansion Phase I (Green Line, 10 stations)',
     '**AIIB**', '**USD 180M**', '**Azerbaijan**', 'Urban rail',
     'Complements EBRD EUR 5B envelope; Azerbaijan consolidating as the most MDB-active South Caucasus sovereign.'],
]
make_table(port_headers, port_rows, col_widths=[3.8, 2.8, 2.5, 1.8, 2.2, 4.5])
section_rule()

# ── 3. Macro Context ──────────────────────────────────────────────────────────
heading('MACRO CONTEXT & INFRASTRUCTURE IMPACT', level=2, colour=NAVY, size=12)

macro_bullets = [
    '**MIDDLE CORRIDOR STRUCTURAL LOCK-IN:** Post-Russia freight rerouting is permanent, not cyclical. Over **USD 12B** committed on the TITR spine in 2025–2026 across rail, road, port, and border infrastructure — creating compounding throughput interdependencies that make individual project returns contingent on corridor-wide completion.',
    '**GREEN TRANSITION AS PRIMARY MDB GATEWAY:** **EBRD** targets 50%+ green financing share; **World Bank** Evolution Roadmap mandates climate integration; **ADB** holds a **USD 100B** climate target through 2030. Green labeling is now baseline eligibility, not a differentiator — sovereigns without NDC-aligned pipelines face growing access constraints.',
    '**PRIVATE CAPITAL MOBILIZATION AS INSTITUTIONAL KPI:** G20 MDB Capital Adequacy Review demands **4:1–7:1 private leverage ratios**. Every major platform vehicle this week — **TIDIP (EUR 5B), EBRD+EU Digital Package (EUR 100M), IFC Highland Fund II (USD 75M target)** — is structured to crowd in commercial capital. Blended finance is the MDB business model, not an option.',
    '**URBAN INFRASTRUCTURE AS NDC ALIGNMENT VEHICLE:** Tramlines, metro expansions, and water/wastewater systems are simultaneously domestically popular, Paris Agreement-eligible, and MDB-bankable. **Türkiye (EUR 191.5M), Azerbaijan (USD 180M metro + EUR 410M water programme), Kazakhstan** are executing this archetype at scale.',
    '**CENTRAL ASIAN DIGITAL INFRASTRUCTURE MAINSTREAMED:** **EDB\'s USD 70M** Uzum fintech loan, **EBRD\'s ~USD 43M** Tajikistan telecoms MOU, and the **EBRD+EU EUR 100M** digital-green package confirm digital infrastructure has crossed from experimental to bankable asset class across the region.',
    '**PARAMETRIC DISASTER RISK FINANCE SCALING:** **ADB\'s USD 150M** catastrophe bonds for **Kyrgyzstan and Tajikistan** (earthquake + precipitation triggers) mark the first systematic shift from project-level resilience to actuarial/fiscal balance sheet management in Central Asia — a replicable sovereign instrument with significant appetite ahead.',
    '**EDB AS PARALLEL FINANCING ARCHITECTURE:** **EDB\'s BBB/Stable** affirmation, Tashkent office opening, and Uzbekistan non-member lending signal deliberate construction of a Western-sanctions-independent development finance tier. Western co-investors must treat EDB-adjacent transactions as secondary sanctions exposure — a risk systematically underweighted in current deal screening.',
]
for b in macro_bullets:
    bullet(b)
section_rule()

# ── 4. Corridor Intelligence ──────────────────────────────────────────────────
heading('CORRIDOR INTELLIGENCE: MIDDLE CORRIDOR CAPITAL STACK', level=2, colour=NAVY, size=12)

corr_headers = ['Node', 'Key Projects', 'MDB Capital Committed', 'Strategic Function']
corr_rows = [
    ['**Eastern Kazakhstan**',
     'KTZ Mointy–Kyzylzhar Rail (322 km) · Almaty Railway Bypass · Saryagash Bypass Road · KTZ Eurobond Station Upgrades',
     '**USD 2.21B** (WB, AIIB, IFC, MIGA, EBRD, ADB)',
     'Primary freight origination hub; greenfield rail closes last eastern gap in TITR land route.'],
    ['**Caspian Crossing**',
     'Alat Port (expanded) · Caspian RoRo/ferry capacity enabled by upstream rail',
     'Committed via upstream investments',
     'Cross-Caspian link; throughput now upstream-constrained (Kazakhstan) and downstream-constrained (Georgia) — both now being addressed.'],
    ['**Azerbaijan**',
     'EBRD EUR 5B envelope · Baku Metro Green Line (USD 180M) · Ganja Water EUR 40M (Phase 1 of EUR 410M) · Karabakh Irrigation USD 436.67M · Sumqayit Wastewater USD 90.4M',
     '**EUR 5B+ envelope; USD 707M+ discrete**',
     'Post-conflict reconstruction node + green urban transition showcase; IsDB hosting amplifies GCC co-investor signal.'],
    ['**Georgia (Land Bridge)**',
     'TC-GATE Rail Electrification + Road Corridors',
     '**USD 729M** (WB + ADB + AIIB)',
     'THE binding throughput constraint on the entire corridor. Electric loco + substation upgrades resolve 3–6M tonne/yr bottleneck. Without TC-GATE, USD 1.41B Kazakhstan rail and USD 6.75B INRAIL capacity are structurally disconnected.'],
    ['**Türkiye / Bosphorus**',
     'INRAIL Istanbul North Rail Crossing (127 km) · Green & Future Cities EUR 191.5M · Arkas Maritime USD 40M · Sustainable OIZ USD 250M',
     '**USD 7.23B+** (WB + 5 MDBs, EBRD, IFC, IsDB, ADB)',
     'Western terminus and global freight gateway; INRAIL\'s 3M→50M tonne/yr uplift is the corridor\'s ultimate value capture point.'],
]
make_table(corr_headers, corr_rows, col_widths=[3.0, 5.2, 3.5, 5.8])
section_rule()

# ── 5. Risk Register ──────────────────────────────────────────────────────────
heading('RISK REGISTER', level=2, colour=NAVY, size=12)

risk_headers = ['Risk', 'Affected Entities', 'Assessment']
risk_rows = [
    ['**KTZ Sovereign Debt Concentration**',
     '**Kazakhstan** MoF, **KTZ**, **EBRD · World Bank · AIIB**',
     'Aggregate sovereign-guaranteed railway exposure exceeds **USD 3B** in 2025–2026 alone. If TITR transit revenues disappoint, contingent liabilities crystallize on an already-stretched sovereign balance sheet. **ELEVATED.**'],
    ['**Azerbaijan Water Sector Absorptive Capacity**',
     '**Azersu · MACP**, **EBRD · OPEC Fund · IsDB**',
     '**EUR 410M+** water programme scope vs. implementing agency procurement capacity creates disbursement bottleneck risk. ADB TA (est. USD 0.5–2M) materially insufficient for transaction volume. Karabakh timeline slippage carries direct political cost. **MODERATE–ELEVATED.**'],
    ['**EDB Sanctions Permeability**',
     '**EDB**, Western co-investors, correspondent banks in EDB-adjacent deals',
     'Russia + Belarus ownership + expanding non-member state footprint (Uzbekistan fintech, Tajikistan logistics) creates OFAC/EU secondary sanctions exposure systematically underpriced in current frameworks. **MODERATE — trending ELEVATED.**'],
]
make_table(risk_headers, risk_rows, hdr_bg=RGBColor(0x7B, 0x24, 0x1C),
           col_widths=[4.0, 4.5, 9.0])
section_rule()

# ── 6. Forward Calendar ───────────────────────────────────────────────────────
heading('FORWARD CALENDAR: 72-HOUR CRITICAL WINDOW', level=2, colour=NAVY, size=12)
body('IsDB Baku Annual Meetings & TIIF Tashkent — June 16–19, 2026', colour=GOLD, bold_phrases=None, size=9.5)

fwd_bullets = [
    'Watch for **IsDB USD 436.67M Karabakh Irrigation** formal signing — converts political signal to disbursement-grade commitment and sets GCC co-investor precedent for post-conflict reconstruction financing.',
    'Monitor **IsDB sukuk framework announcement** and **Alat FEZ GCC co-investment MOUs** — would materially expand Gulf capital\'s structural role in Azerbaijan\'s reconstruction corridor.',
    'At **TIIF Tashkent**, track whether **NDB makes its first Uzbekistan engagement** public — NDB entry would signal BRICS-aligned capital competing directly with Western MDB platforms in Central Asia.',
    'Watch for **ADB Uzbekistan energy sector announcement** at TIIF — likely a large-scale renewable or grid commitment extending the Bash 2 Wind (300 MW) template.',
    'Monitor **IFC Highland Fund II LP close progress** — whether **NFRK (Kazakhstan)** or **UzFund (Uzbekistan)** anchor commitments are confirmed determines the fund\'s private capital mobilization credibility.',
    'Apply **MOU-to-disbursement conversion scrutiny** across all TIIF announcements — the execution gap between headline figures and first-disbursement timelines is the primary risk in the Uzbekistan market.',
]
for b in fwd_bullets:
    bullet(b)
section_rule()

# ── 7. Opportunity Flags ──────────────────────────────────────────────────────
heading('OPPORTUNITY FLAGS', level=2, colour=GREEN_H, size=12)

opp_bullets = [
    '**Tajikistan as underpriced infrastructure entry point:** Five distinct MDB transactions this week, aggregating **>USD 900M** in commitments against a **USD 12B GDP** economy — the highest MDB capital intensity ratio in the region. **AIIB\'s USD 800M pipeline** remains underpopulated on the private co-investment side; first-mover positioning across energy, logistics, and digital is available now at pre-competition pricing.',
    '**TIIF Tashkent as compressed LP access window:** **8,300+ participants**, ministerial-level Uzbek counterparts, and eight MDB heads co-located for 72 hours is the highest-density Central Asian capital introduction environment of 2026. Fund managers targeting **NFRK or UzFund LP commitments** should deploy senior relationship capital this week — the **IFC Highland Fund II** co-investment narrative provides a credible institutional anchor.',
]
for b in opp_bullets:
    bullet(b)

# ── Footer ────────────────────────────────────────────────────────────────────
section_rule()
p_foot = doc.add_paragraph()
p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_foot = p_foot.add_run(
    'CLASSIFICATION: Senior Distribution — C-Suite & Board Advisory  |  Prepared: June 15, 2026  |  Not for public release or redistribution'
)
r_foot.italic = True
r_foot.font.size = Pt(8)
r_foot.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
r_foot.font.name = 'Calibri'

# ── Save ──────────────────────────────────────────────────────────────────────
output_path = '/home/user/claude-code/MDB_IFI_Intelligence_Brief_June15_2026.docx'
doc.save(output_path)
print(f'Saved: {output_path}')
